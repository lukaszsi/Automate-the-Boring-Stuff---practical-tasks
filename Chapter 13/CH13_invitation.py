#! python3
import docx
import os

guestsFile = open('guests.txt')
text = guestsFile.read()
guestsList = text.split('\n')
os.chdir('C:\\Python')
doc = docx.Document('invitation.docx')
for guest in guestsList:
    invitaionDictionary = {'It would be a pleasure to have a company of': 'Invitation',
                           guest: 'Guests',
                           'at 111 Memory Lane on the Evening of': 'Invitation',
                           'April 1st': 'Guests',
                           'at 7\'o clock': 'Invitation'}
    for k, v in invitaionDictionary.items():
        doc.add_paragraph(k, v)
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.save('invitationDone.docx')
